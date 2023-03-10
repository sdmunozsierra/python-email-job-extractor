import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

def format_education(doc, education):

    for edu in education:
        #p = doc.add_paragraph()
        #p.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
        text = f"Bachelor of Science in {edu.major} and Minor in {edu.minor} - B.S.CS"
        h = doc.add_heading(text, level=3)
        h.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        p.add_run(f"\n{edu.school_name} at {edu.location}:\n").bold = True
        p.add_run(f"GPA: {edu.gpa}").italic = True

        p.add_run(f"\nCoursework: ").bold = True
        for i in edu.coursework:
            p.add_run(", ".join([str(x) for x in i]))

        p.add_run(f"\nOrganizations: ").bold = True
        for i in edu.organizations:
            p.add_run(", ".join([str(x) for x in i]))

        p.add_run(f"\nResearch: ").bold = True
        for i in edu.research:
            p.add_run(", ".join([str(x) for x in i]))

