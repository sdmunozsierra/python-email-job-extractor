import docx
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from src.person_builder import PersonBuilder
from src.education_builder import EducationFactory
from src.cert_builder import CertFactory
from src.format_experience import format_experience, format_experience_skills
from personal_info import sergio_david_munoz_sierra as sdms


utep_education = EducationFactory.create_education(
    "Computer Science",
    "Math",
    "The University of Texas at El Paso",
    "Dic 2018",
    "El Paso, TX",
    2.4,
    ["Artificial Intelligence", "Numerical Analysis", "Computer Networks", "Software Construction", "Secure Web Systems"],
    ["Google Ignite CS Program", "Miner's Cyber Security Club (MSSC)"],
    ["Dr. Omar Baddredin - Crowd-sourcing Road Topology and Driving Patterns using Smartphone's Sensors"],
    ["Research with Dr. Omar Baddredin - UTD - 4th Place "]
)

certs = CertFactory.from_file("certs.txt")
#for cert in certs:
  # print(cert)

sergio_builder = PersonBuilder()\
    .add_experience(sdms.job0)\
    .add_experience(sdms.job1)\
    .add_experience(sdms.job2)\
    .add_experience(sdms.job3)\
    .add_experience(sdms.job4)\
    .add_experience(sdms.job5)\
    .add_experience(sdms.job6)\
    .add_experience(sdms.job7)\
    .set_skills(
        ["postman", "REST", "sonarqube", "swagger-api",
                "aws-aurora","mongo-db", "graphql", "redis", "sql",
                "java-8+","java-android", "mvn",
                "js-node", "js-vue", "js-nuxt", "js-express",
                "python2", "python3", "python-scrapy", "python-flask", "python-pandas",
                "spring-boot", "spring-security", "spring-profiles", "spring-data",
                "bash", "kafka", "vim", "yarn","solidity",
                "html", "css", "php", "jinja",
                "git", "github-actions", "gitlab-cicd",
                "aws-amplify", "aws-iam", "aws-serverless", "aws-vpc", "aws-codepipeline", "aws-cloudformation",
                "ansible", "consul", "kubernetes-training",
                "docker", "docker-compose", "docker-swarm", "docker-images",
                "elastic-stack", "elastic-search", "log-stash", "kibana",
                "arch-linux", "centos", "debian", "mint", "ubuntu-flavors", "proxmox",
                "english-fluent", "spanish-fluent", "french-wp"])\
    .add_activity("1st Place - 2019 RESET Hackathon (Blockchain)")\
    .add_activity("2nd Place - 2021 Verac Hackathon (Android)")\
    .add_activity("3rd Place - 2018 UTD Research Symposium (Android)")\
    .add_activity("4th Place - 2020 SANS Institute CTF (Pen Testing)")\
    .set_certs(certs)
sergio = sergio_builder.build()
#print(sergio)
#for e in sergio.experience:
    #print(e)

def build_resume(person):
    # Create a Word document object
    doc = docx.Document()

    # Add a header with the name and job title
    header = doc.add_heading(person.name, level=1)
    header.add_run("\n" + person.job_title).bold = True
    doc.add_paragraph()

    # Add a section for experience
    doc.add_heading("Experience", level=2)
    format_experience(doc, person.experience)
    #for experience_item in person.experience:
        #doc.add_paragraph(str(experience_item))
    doc.add_paragraph()

    # Add a section for education
    #doc.add_heading("Education", level=2)
    #format_education(doc, person.education)
    #doc.add_paragraph()

    # Add a section for skills
    doc.add_heading("Skills", level=2)
    doc.add_heading("Experience Skills", level=3)
    format_experience_skills(doc, person.experience)

    doc.add_heading("Related Skills (A-z)", level=3)
    arr = [x for x in person.skills]
    arr.sort()
    p = doc.add_paragraph()
    #p.alignment = WD_STYLE_TYPE.JUSTIFY
    p.add_run(", ".join([str(x) for x in arr]))
    doc.add_paragraph()

    # Add a section for activities
    doc.add_heading("Activities", level=2)
    p = doc.add_paragraph()
    #p.alignment = WD_STYLE_TYPE.JUSTIFY
    p.add_run(", ".join([str(x) for x in person.activities]))
    #for activity in person.activities:
        #doc.add_paragraph(activity)

    # Add a section for certs
    doc.add_heading("Certs", level=2)
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Name'
    hdr_cells[1].text = 'Issue'
    hdr_cells[2].text = 'Date'
    for cert in person.certifications:
        row_cells = table.add_row().cells
        row_cells[0].text = cert.title
        row_cells[1].text = cert.issuer
        row_cells[2].text = cert.completion_date

    # Iterate through all styles and set their font to the new default font
    #for style in doc.styles:
        #if style.type == WD_STYLE_TYPE.PARAGRAPH:
            #font = style.font
            #font.name = 'Arial'
            #font.size = Pt(12)
    return doc

# Example usage
sergio.name = "Sergio David Munoz Sierra"
sergio.job_title = "Future Master of Science in Analytics"

resume = build_resume(sergio)

# Save the Word document to a file
resume.save("resume.docx")
