from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_border(cell, border_pos, border_width, border_color="auto"):
    """ Set border of a cell for the timeline """
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()

    # Create the cell borders as needed
    borders = OxmlElement('w:tcBorders')

    border = OxmlElement(f'w:{border_pos}')
    border.set(qn('w:sz'), str(border_width))
    border.set(qn('w:val'), 'single')
    border.set(qn('w:color'), border_color)
    borders.append(border)

    tcPr.append(borders)

def add_timeline_entry(table, position, col_idx):
    """ Add timeline entry (title and summary) to the table """
    # Determine the row index based on column index
    is_even_col = col_idx % 2 == 0
    title_row_idx = 1 if is_even_col else 2
    summary_row_idx = 0 if is_even_col else 3

    # Title cell
    title_cell = table.cell(title_row_idx, col_idx)
    title_paragraph = title_cell.paragraphs[0]
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(f"{position['title']} ({position['start']} - {position['end']})")
    title_run.bold = True
    title_run.font.size = Pt(12)

    # Summary cell
    summary_cell = table.cell(summary_row_idx, col_idx)
    summary_paragraph = summary_cell.paragraphs[0]
    summary_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary_run = summary_paragraph.add_run(position['summary'])
    summary_run.font.size = Pt(10)

def add_alternating_timeline(doc, positions):
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Professional Experience Timeline')
    run.bold = True
    run.font.size = Pt(14)

    # Create a table for the timeline
    table = doc.add_table(rows=4, cols=len(positions) * 2)

    # Add timeline entries
    for i, position in enumerate(positions):
        add_timeline_entry(table, position, i)


# Example positions data
positions = [
    {'title': 'Software Developer', 'start': '2015', 'end': '2017', 'summary': 'Developed key software solutions...'},
    {'title': 'Senior Developer', 'start': '2017', 'end': '2019', 'summary': 'Led a team of developers...'},
    {'title': 'Project Manager', 'start': '2019', 'end': '2021', 'summary': 'Managed multiple software projects...'},
    {'title': 'Senior Project Manager', 'start': '2021', 'end': '2023', 'summary': 'Oversaw large-scale projects...'},
    # Additional positions as needed
]

# Create a new Word document
doc = Document()
add_alternating_timeline(doc, positions)

# Save the document
doc.save('professional_timeline_cv.docx')