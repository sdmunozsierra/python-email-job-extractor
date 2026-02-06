"""Professional document styling for polished resume DOCX output.

Provides reusable helpers for typography, colors, margins, section dividers,
and custom paragraph/table styles used throughout the resume builder.
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy


# ---------------------------------------------------------------------------
# Color palette — professional dark navy / slate / accent blue
# ---------------------------------------------------------------------------

class Colors:
    """Central color palette for the resume."""
    PRIMARY = RGBColor(0x1B, 0x2A, 0x4A)       # Dark navy — name, section headings
    SECONDARY = RGBColor(0x2C, 0x3E, 0x6B)     # Medium navy — subtitles
    ACCENT = RGBColor(0x3B, 0x7D, 0xD8)        # Accent blue — links, highlights
    TEXT = RGBColor(0x2D, 0x2D, 0x2D)           # Near-black — body text
    TEXT_LIGHT = RGBColor(0x55, 0x55, 0x55)     # Medium gray — secondary text
    TEXT_MUTED = RGBColor(0x77, 0x77, 0x77)     # Muted gray — dates, captions
    DIVIDER = RGBColor(0x3B, 0x7D, 0xD8)       # Accent blue — section dividers
    TABLE_HEADER_BG = "1B2A4A"                  # Dark navy hex (for XML)
    TABLE_ALT_BG = "F0F4FA"                     # Very light blue (for XML)
    TABLE_BORDER = "CCCCCC"                     # Light gray border
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)


# ---------------------------------------------------------------------------
# Font families
# ---------------------------------------------------------------------------

class Fonts:
    """Font families used in the resume."""
    HEADING = "Calibri"
    BODY = "Calibri"
    MONO = "Consolas"


# ---------------------------------------------------------------------------
# Document setup
# ---------------------------------------------------------------------------

def create_styled_document():
    """Create a new Document with professional margins and default styles."""
    doc = Document()

    # Set narrow margins for maximum content area
    for section in doc.sections:
        section.top_margin = Cm(1.2)
        section.bottom_margin = Cm(1.2)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    # Set default font for the document
    style = doc.styles["Normal"]
    font = style.font
    font.name = Fonts.BODY
    font.size = Pt(10)
    font.color.rgb = Colors.TEXT

    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(2)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

    return doc


# ---------------------------------------------------------------------------
# Typography helpers
# ---------------------------------------------------------------------------

def add_name_heading(doc, name):
    """Add the candidate's name as a large, bold heading."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(name.upper())
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = Colors.PRIMARY
    run.font.name = Fonts.HEADING
    return p


def add_job_title(doc, title):
    """Add the job title / tagline centered below the name."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    run.font.size = Pt(12)
    run.font.color.rgb = Colors.SECONDARY
    run.font.name = Fonts.HEADING
    run.italic = True
    return p


def add_contact_line(doc, parts):
    """Add a centered contact info line with pipe separators."""
    if not parts:
        return None
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)

    for i, part in enumerate(parts):
        if i > 0:
            sep = p.add_run("  |  ")
            sep.font.size = Pt(9)
            sep.font.color.rgb = Colors.TEXT_MUTED
            sep.font.name = Fonts.BODY
        run = p.add_run(part)
        run.font.size = Pt(9)
        run.font.color.rgb = Colors.TEXT_LIGHT
        run.font.name = Fonts.BODY
    return p


def add_summary(doc, text):
    """Add the professional summary as an italicized block."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.right_indent = Inches(0.15)
    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.color.rgb = Colors.TEXT
    run.font.name = Fonts.BODY
    run.italic = True
    return p


# ---------------------------------------------------------------------------
# Section dividers & headings
# ---------------------------------------------------------------------------

def add_section_heading(doc, title):
    """Add a section heading with a colored bottom border line.

    Returns the paragraph so callers can adjust spacing if needed.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.keep_with_next = True

    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = Colors.PRIMARY
    run.font.name = Fonts.HEADING

    # Add a bottom border using XML (a thin colored line under the paragraph)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="{Colors.DIVIDER.hex()}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    return p


def add_subsection_heading(doc, title):
    """Add a lighter subsection heading (e.g. skill categories)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True

    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = Colors.SECONDARY
    run.font.name = Fonts.HEADING
    return p


# ---------------------------------------------------------------------------
# RGBColor helper — hex string for XML
# ---------------------------------------------------------------------------

def _rgb_hex(color):
    """Return a 6-char hex string from an RGBColor."""
    if hasattr(color, 'hex'):
        return color.hex()
    return f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"


# Monkey-patch RGBColor.hex if not present
if not hasattr(RGBColor, 'hex'):
    RGBColor.hex = lambda self: f"{self[0]:02X}{self[1]:02X}{self[2]:02X}"


# ---------------------------------------------------------------------------
# Experience entry helpers
# ---------------------------------------------------------------------------

def add_experience_header(doc, role, company, location, date_range):
    """Add an experience entry header with role/company on the left and dates on the right.

    Uses a 1-row, 2-column invisible table to align content.
    """
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Remove all table borders
    _remove_table_borders(table)

    # Set table width to full page
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>') 
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    tblPr.append(tblW)

    # Left cell: Role + Company
    left_cell = table.rows[0].cells[0]
    left_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    lp = left_cell.paragraphs[0]
    lp.paragraph_format.space_before = Pt(0)
    lp.paragraph_format.space_after = Pt(0)

    role_run = lp.add_run(role)
    role_run.bold = True
    role_run.font.size = Pt(10.5)
    role_run.font.color.rgb = Colors.PRIMARY
    role_run.font.name = Fonts.HEADING

    sep_run = lp.add_run("  —  ")
    sep_run.font.size = Pt(10)
    sep_run.font.color.rgb = Colors.TEXT_MUTED
    sep_run.font.name = Fonts.BODY

    co_run = lp.add_run(company)
    co_run.font.size = Pt(10)
    co_run.font.color.rgb = Colors.SECONDARY
    co_run.font.name = Fonts.BODY
    co_run.italic = True

    if location:
        loc_run = lp.add_run(f"  ({location})")
        loc_run.font.size = Pt(9)
        loc_run.font.color.rgb = Colors.TEXT_MUTED
        loc_run.font.name = Fonts.BODY

    # Right cell: Date range
    right_cell = table.rows[0].cells[1]
    right_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    rp = right_cell.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.space_before = Pt(0)
    rp.paragraph_format.space_after = Pt(0)

    date_run = rp.add_run(date_range)
    date_run.font.size = Pt(9)
    date_run.font.color.rgb = Colors.TEXT_MUTED
    date_run.font.name = Fonts.BODY
    date_run.italic = True

    # Set column widths (roughly 70/30 split)
    for cell in table.rows[0].cells:
        _set_cell_margins(cell, top=0, bottom=0, start=0, end=0)

    return table


def add_experience_description(doc, text):
    """Add a brief role description in italic."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.15)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = Colors.TEXT_LIGHT
    run.font.name = Fonts.BODY
    return p


def add_bullet_point(doc, text):
    """Add a single bullet point with proper formatting."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.first_line_indent = Inches(-0.15)

    # Bullet character
    bullet_run = p.add_run("\u2022  ")
    bullet_run.font.size = Pt(9.5)
    bullet_run.font.color.rgb = Colors.ACCENT
    bullet_run.font.name = Fonts.BODY

    # Bullet text
    text_run = p.add_run(text)
    text_run.font.size = Pt(9.5)
    text_run.font.color.rgb = Colors.TEXT
    text_run.font.name = Fonts.BODY
    return p


def add_tech_tags(doc, technologies):
    """Add a technologies line with bold label and comma-separated values."""
    if not technologies:
        return None
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.15)

    label_run = p.add_run("Technologies: ")
    label_run.bold = True
    label_run.font.size = Pt(9)
    label_run.font.color.rgb = Colors.SECONDARY
    label_run.font.name = Fonts.BODY

    tech_run = p.add_run(", ".join(technologies))
    tech_run.font.size = Pt(9)
    tech_run.font.color.rgb = Colors.TEXT_LIGHT
    tech_run.font.name = Fonts.BODY
    tech_run.italic = True
    return p


# ---------------------------------------------------------------------------
# Table helpers
# ---------------------------------------------------------------------------

def create_styled_table(doc, headers, rows, col_widths=None):
    """Create a professionally styled table with header row and alternating colors.

    Args:
        doc: Document instance
        headers: list of header strings
        rows: list of lists (each inner list is a row of cell values)
        col_widths: optional list of Inches values for column widths
    """
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Style header row
    hdr = table.rows[0]
    for i, header_text in enumerate(headers):
        cell = hdr.cells[i]
        _shade_cell(cell, Colors.TABLE_HEADER_BG)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(header_text)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = Colors.WHITE
        run.font.name = Fonts.HEADING
        _set_cell_margins(cell, top=40, bottom=40, start=80, end=80)

    # Style data rows with alternating background
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        for col_idx, value in enumerate(row_data):
            cell = row.cells[col_idx]
            if row_idx % 2 == 1:
                _shade_cell(cell, Colors.TABLE_ALT_BG)
            p = cell.paragraphs[0]
            run = p.add_run(str(value) if value else "")
            run.font.size = Pt(9)
            run.font.color.rgb = Colors.TEXT
            run.font.name = Fonts.BODY
            _set_cell_margins(cell, top=30, bottom=30, start=80, end=80)

    # Set thin borders
    _set_table_borders(table, Colors.TABLE_BORDER)

    # Set column widths if provided
    if col_widths and len(col_widths) == len(headers):
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = width

    return table


def create_skills_table(doc, skills_by_category):
    """Create a compact two-column skills table grouped by category.

    Args:
        skills_by_category: dict mapping category name to list of skill strings
    """
    if not skills_by_category:
        return None

    rows_data = []
    for category, skill_list in skills_by_category.items():
        rows_data.append((category, ", ".join(skill_list)))

    table = doc.add_table(rows=len(rows_data), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Set table width to 100%
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    tblPr.append(tblW)

    for row_idx, (category, skills_text) in enumerate(rows_data):
        row = table.rows[row_idx]

        # Category cell (left)
        cat_cell = row.cells[0]
        if row_idx % 2 == 0:
            _shade_cell(cat_cell, Colors.TABLE_ALT_BG)
        p = cat_cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(category)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = Colors.PRIMARY
        run.font.name = Fonts.HEADING
        _set_cell_margins(cat_cell, top=30, bottom=30, start=80, end=40)
        cat_cell.width = Inches(1.8)

        # Skills cell (right)
        skill_cell = row.cells[1]
        if row_idx % 2 == 0:
            _shade_cell(skill_cell, Colors.TABLE_ALT_BG)
        p = skill_cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(skills_text)
        run.font.size = Pt(9)
        run.font.color.rgb = Colors.TEXT
        run.font.name = Fonts.BODY
        _set_cell_margins(skill_cell, top=30, bottom=30, start=40, end=80)

    # Light borders
    _set_table_borders(table, Colors.TABLE_BORDER, inner=True)

    return table


# ---------------------------------------------------------------------------
# Education entry helper
# ---------------------------------------------------------------------------

def add_education_entry(doc, title, institution, location, date_range,
                        gpa=None, honors=None, coursework=None):
    """Add a formatted education entry."""
    # Header row with title on left, dates on right
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _remove_table_borders(table)

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    tblPr.append(tblW)

    left_cell = table.rows[0].cells[0]
    left_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    lp = left_cell.paragraphs[0]
    lp.paragraph_format.space_before = Pt(0)
    lp.paragraph_format.space_after = Pt(0)

    title_run = lp.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(10.5)
    title_run.font.color.rgb = Colors.PRIMARY
    title_run.font.name = Fonts.HEADING

    # Institution on next line
    inst_p = left_cell.add_paragraph()
    inst_p.paragraph_format.space_before = Pt(0)
    inst_p.paragraph_format.space_after = Pt(0)
    inst_run = inst_p.add_run(institution)
    inst_run.font.size = Pt(9.5)
    inst_run.font.color.rgb = Colors.SECONDARY
    inst_run.font.name = Fonts.BODY
    inst_run.italic = True

    if location:
        loc_run = inst_p.add_run(f"  —  {location}")
        loc_run.font.size = Pt(9)
        loc_run.font.color.rgb = Colors.TEXT_MUTED
        loc_run.font.name = Fonts.BODY

    right_cell = table.rows[0].cells[1]
    right_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    rp = right_cell.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.space_before = Pt(0)
    rp.paragraph_format.space_after = Pt(0)

    if date_range:
        date_run = rp.add_run(date_range)
        date_run.font.size = Pt(9)
        date_run.font.color.rgb = Colors.TEXT_MUTED
        date_run.font.name = Fonts.BODY
        date_run.italic = True

    for cell in table.rows[0].cells:
        _set_cell_margins(cell, top=0, bottom=0, start=0, end=0)

    # Details below
    details = []
    if gpa:
        details.append(("GPA", str(gpa)))
    if honors:
        details.append(("Honors", honors))

    if details:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.left_indent = Inches(0.15)
        for i, (label, value) in enumerate(details):
            if i > 0:
                p.add_run("  |  ").font.color.rgb = Colors.TEXT_MUTED
            lbl = p.add_run(f"{label}: ")
            lbl.bold = True
            lbl.font.size = Pt(9)
            lbl.font.color.rgb = Colors.SECONDARY
            lbl.font.name = Fonts.BODY
            val = p.add_run(value)
            val.font.size = Pt(9)
            val.font.color.rgb = Colors.TEXT
            val.font.name = Fonts.BODY

    if coursework:
        flat = []
        for item in coursework:
            if isinstance(item, list):
                flat.extend(item)
            else:
                flat.append(item)
        if flat:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.left_indent = Inches(0.15)
            lbl = p.add_run("Relevant Coursework: ")
            lbl.bold = True
            lbl.font.size = Pt(9)
            lbl.font.color.rgb = Colors.SECONDARY
            lbl.font.name = Fonts.BODY
            val = p.add_run(", ".join(str(c) for c in flat))
            val.font.size = Pt(9)
            val.font.color.rgb = Colors.TEXT_LIGHT
            val.font.name = Fonts.BODY
            val.italic = True


# ---------------------------------------------------------------------------
# Project entry helper
# ---------------------------------------------------------------------------

def add_project_entry(doc, name, description, url=None, bullets=None, technologies=None):
    """Add a formatted project entry."""
    # Project name
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.keep_with_next = True

    name_run = p.add_run(name)
    name_run.bold = True
    name_run.font.size = Pt(10)
    name_run.font.color.rgb = Colors.PRIMARY
    name_run.font.name = Fonts.HEADING

    if url:
        url_run = p.add_run(f"  —  {url}")
        url_run.font.size = Pt(9)
        url_run.font.color.rgb = Colors.ACCENT
        url_run.font.name = Fonts.BODY

    # Description
    if description:
        dp = doc.add_paragraph()
        dp.paragraph_format.space_before = Pt(0)
        dp.paragraph_format.space_after = Pt(2)
        dp.paragraph_format.left_indent = Inches(0.15)
        desc_run = dp.add_run(description)
        desc_run.italic = True
        desc_run.font.size = Pt(9.5)
        desc_run.font.color.rgb = Colors.TEXT_LIGHT
        desc_run.font.name = Fonts.BODY

    # Bullets
    if bullets:
        for bullet in bullets:
            add_bullet_point(doc, bullet)

    # Technologies
    if technologies:
        add_tech_tags(doc, technologies)


# ---------------------------------------------------------------------------
# Certification table helper
# ---------------------------------------------------------------------------

def add_certifications_table(doc, certifications):
    """Add a styled certifications table."""
    headers = ["Certification", "Issuer", "Date"]
    rows = []
    for cert in certifications:
        rows.append([
            cert.title or "",
            cert.issuer or "",
            cert.completion_date or "",
        ])

    return create_styled_table(
        doc, headers, rows,
        col_widths=[Inches(3.5), Inches(2.0), Inches(1.0)]
    )


# ---------------------------------------------------------------------------
# Activities / Awards helpers
# ---------------------------------------------------------------------------

def add_inline_list(doc, items, label=None):
    """Add a comma-separated inline list with optional bold label."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.15)

    if label:
        lbl = p.add_run(f"{label}: ")
        lbl.bold = True
        lbl.font.size = Pt(9.5)
        lbl.font.color.rgb = Colors.SECONDARY
        lbl.font.name = Fonts.BODY

    val = p.add_run(", ".join(str(x) for x in items))
    val.font.size = Pt(9.5)
    val.font.color.rgb = Colors.TEXT
    val.font.name = Fonts.BODY
    return p


def add_activity_bullets(doc, activities):
    """Add activities as bullet points."""
    for activity in activities:
        add_bullet_point(doc, str(activity))


# ---------------------------------------------------------------------------
# Preferences helper
# ---------------------------------------------------------------------------

def add_preferences_section(doc, prefs):
    """Add a compact preferences section."""
    fields = [
        ("desired_roles", "Desired Roles"),
        ("industries", "Industries"),
        ("locations", "Preferred Locations"),
        ("engagement_types", "Engagement Types"),
    ]

    for key, label in fields:
        val = prefs.get(key)
        if val:
            if isinstance(val, list):
                add_inline_list(doc, val, label=label)
            else:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.left_indent = Inches(0.15)
                lbl = p.add_run(f"{label}: ")
                lbl.bold = True
                lbl.font.size = Pt(9.5)
                lbl.font.color.rgb = Colors.SECONDARY
                lbl.font.name = Fonts.BODY
                v = p.add_run(str(val))
                v.font.size = Pt(9.5)
                v.font.color.rgb = Colors.TEXT
                v.font.name = Fonts.BODY

    rp = prefs.get("remote_preference")
    if rp:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.15)
        lbl = p.add_run("Remote Preference: ")
        lbl.bold = True
        lbl.font.size = Pt(9.5)
        lbl.font.color.rgb = Colors.SECONDARY
        lbl.font.name = Fonts.BODY
        v = p.add_run(rp.capitalize())
        v.font.size = Pt(9.5)
        v.font.color.rgb = Colors.TEXT
        v.font.name = Fonts.BODY


# ---------------------------------------------------------------------------
# Internal XML helpers for table styling
# ---------------------------------------------------------------------------

def _shade_cell(cell, color_hex):
    """Apply a background color to a table cell via XML."""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_margins(cell, top=0, bottom=0, start=0, end=0):
    """Set cell margins in twentieths-of-a-point via XML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:start w:w="{start}" w:type="dxa"/>'
        f'  <w:end w:w="{end}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcMar_existing = tcPr.find(qn('w:tcMar'))
    if tcMar_existing is not None:
        tcPr.remove(tcMar_existing)
    tcPr.append(tcMar)


def _remove_table_borders(table):
    """Remove all borders from a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)


def _set_table_borders(table, color_hex, inner=True):
    """Set thin borders on a table."""
    border_parts = [
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>',
        f'<w:left w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>',
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>',
        f'<w:right w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>',
    ]
    if inner:
        border_parts.extend([
            f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>',
            f'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="{color_hex}"/>',
        ])

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders_xml = f'<w:tblBorders {nsdecls("w")}>{"".join(border_parts)}</w:tblBorders>'
    borders = parse_xml(borders_xml)

    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)


def add_spacer(doc, pts=4):
    """Add a small vertical spacer paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    pf = p.paragraph_format
    pf.line_spacing = Pt(pts)
    # Make the run tiny so it's essentially invisible
    run = p.add_run("")
    run.font.size = Pt(1)
    return p
