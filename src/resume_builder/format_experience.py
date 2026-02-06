"""Format experience entries for Word documents with professional styling.

Supports both:
  - Legacy format: experience with nested projects (actions + skills per project)
  - JSON schema format: experience with flat achievements + technologies lists
"""

from resume_builder.document_styles import (
    add_experience_header,
    add_experience_description,
    add_bullet_point,
    add_tech_tags,
    Colors,
    Fonts,
)
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def format_experience(doc, experience):
    """Add experience entries to *doc* with professional formatting."""
    for exp in experience:
        # Date range
        date_str = getattr(exp, "date_range", None) or exp.dates or ""

        # Header: Role — Company (Location)  |  Dates
        add_experience_header(doc, exp.role, exp.company_name, exp.location, date_str)

        # Description (new schema)
        if getattr(exp, "description", None):
            add_experience_description(doc, exp.description)

        # Legacy project-based format
        if exp.projects:
            for proj in exp.projects:
                items = proj if isinstance(proj, list) else [proj]
                for j in items:
                    # Project sub-header
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(3)
                    p.paragraph_format.space_after = Pt(1)
                    p.paragraph_format.left_indent = Inches(0.15)
                    p.paragraph_format.keep_with_next = True

                    name_run = p.add_run(j.name)
                    name_run.bold = True
                    name_run.font.size = Pt(9.5)
                    name_run.font.color.rgb = Colors.SECONDARY
                    name_run.font.name = Fonts.HEADING

                    if j.duration:
                        dur_run = p.add_run(f"  ({j.duration})")
                        dur_run.font.size = Pt(9)
                        dur_run.font.color.rgb = Colors.TEXT_MUTED
                        dur_run.font.name = Fonts.BODY

                    # Project description
                    if j.description:
                        dp = doc.add_paragraph()
                        dp.paragraph_format.space_before = Pt(0)
                        dp.paragraph_format.space_after = Pt(1)
                        dp.paragraph_format.left_indent = Inches(0.25)
                        desc_run = dp.add_run(j.description)
                        desc_run.italic = True
                        desc_run.font.size = Pt(9)
                        desc_run.font.color.rgb = Colors.TEXT_LIGHT
                        desc_run.font.name = Fonts.BODY

                    # Action bullets
                    for a in j.actions:
                        add_bullet_point(doc, a)

        # Flat achievements (new schema)
        if getattr(exp, "achievements", None):
            for achievement in exp.achievements:
                add_bullet_point(doc, achievement)

        # Technologies list (new schema)
        if getattr(exp, "technologies", None):
            add_tech_tags(doc, exp.technologies)


def format_experience_skills(doc, experience):
    """Add per-experience skill summaries to *doc*."""
    for exp in experience:
        # Legacy project-based skills
        if exp.projects:
            for proj in exp.projects:
                items = proj if isinstance(proj, list) else [proj]
                for j in items:
                    if j.skills:
                        p = doc.add_paragraph()
                        p.paragraph_format.space_before = Pt(1)
                        p.paragraph_format.space_after = Pt(1)
                        p.paragraph_format.left_indent = Inches(0.15)

                        label = p.add_run(f"{j.name}: ")
                        label.bold = True
                        label.font.size = Pt(9)
                        label.font.color.rgb = Colors.SECONDARY
                        label.font.name = Fonts.HEADING

                        val = p.add_run(", ".join(str(x) for x in j.skills))
                        val.font.size = Pt(9)
                        val.font.color.rgb = Colors.TEXT
                        val.font.name = Fonts.BODY

        # Flat technologies (new schema)
        if getattr(exp, "technologies", None) and not exp.projects:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.left_indent = Inches(0.15)

            label = p.add_run(f"{exp.role} @ {exp.company_name}: ")
            label.bold = True
            label.font.size = Pt(9)
            label.font.color.rgb = Colors.SECONDARY
            label.font.name = Fonts.HEADING

            val = p.add_run(", ".join(exp.technologies))
            val.font.size = Pt(9)
            val.font.color.rgb = Colors.TEXT
            val.font.name = Fonts.BODY
