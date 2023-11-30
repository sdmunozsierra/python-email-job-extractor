from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_border(cell, border_pos, border_width):
    """ Set border of a cell for the timeline """
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()

    # Create the cell borders as needed
    borders = OxmlElement('w:tcBorders')

    border = OxmlElement(f'w:{border_pos}')
    border.set(qn('w:sz'), str(border_width))
    border.set(qn('w:val'), 'single')
    border.set(qn('w:color'), 'auto')
    borders.append(border)

    tcPr.append(borders)

def add_timeline_section(table, positions, start_idx, end_idx, offset=0, is_top=True):
    """ Add timeline section (top or bottom) """
    for i in range(start_idx, end_idx):
        col_idx = 2 * (i - start_idx) + offset
        position = positions[i]

        # Position title cell
        title_cell = table.cell(0, col_idx)
        title_p = title_cell.paragraphs[0]
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_p.add_run(f"{position['title']} ({position['start']} - {position['end']})")
        run.bold = True

        # Merge with next cell for spacing, if not the last cell
        if col_idx + 1 < len(table.columns):
            title_cell.merge(table.cell(0, col_idx + 1))

        # Add timeline border to cell
        border_pos = 'top' if is_top else 'bottom'
        set_cell_border(title_cell, border_pos, 4)  # Adjust border width as needed

def add_alternating_timeline(doc, positions):
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Professional Experience Timeline')
    run.bold = True
    run.font.size = Pt(14)

    # Calculate the number of positions for the top and bottom tables
    half_len = len(positions) // 2

    # Create top table for the first half of positions
    top_table = doc.add_table(rows=1, cols=2 * half_len)
    add_timeline_section(top_table, positions, 0, half_len, is_top=True)

    # Create bottom table for the second half of positions, offset by one cell
    bottom_table = doc.add_table(rows=1, cols=2 * (len(positions) - half_len) + 1)
    add_timeline_section(bottom_table, positions, half_len, len(positions), offset=1, is_top=False)

# Example positions data
positions = [
    {'title': 'Software Developer', 'start': '2015', 'end': '2017', 'summary': 'Developed key software solutions...'},
    {'title': 'Senior Developer', 'start': '2017', 'end': '2019', 'summary': 'Led a team of developers...'},
    {'title': 'Project Manager', 'start': '2019', 'end': '2021', 'summary': 'Managed multiple software projects...'},
    {'title': 'Senior Project Manager', 'start': '2021', 'end': '2023', 'summary': 'Oversaw large-scale projects...'},
    # Additional positions as needed
]

# Create a new Word document
doc = Document()
add_alternating_timeline(doc, positions)

# Save the document
doc.save('alternating_timeline_cv.docx')

