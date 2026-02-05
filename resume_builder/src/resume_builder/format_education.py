"""Word-doc education formatting helpers."""

from __future__ import annotations

from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from docx import Document

from .education_builder import Education


def add_education_section(doc: "Document", education_list: list) -> None:
    """Add an Education section to a Word document.

    Args:
        doc: python-docx Document instance.
        education_list: List of :class:`Education` objects.
    """
    if not education_list:
        return

    doc.add_heading("Education", level=1)

    for edu in education_list:
        degree_line = edu.degree
        if edu.field:
            degree_line += f" in {edu.field}"
        if edu.school_name:
            degree_line += f" -- {edu.school_name}"
        doc.add_heading(degree_line, level=2)

        meta_parts = []
        if edu.location:
            meta_parts.append(edu.location)
        if edu.dates:
            meta_parts.append(edu.dates)
        if meta_parts:
            doc.add_paragraph(" | ".join(meta_parts)).italic = True

        if edu.gpa:
            doc.add_paragraph(f"GPA: {edu.gpa}")

        if edu.honors:
            doc.add_paragraph(f"Honors: {edu.honors}")

        if edu.coursework:
            doc.add_paragraph(f"Relevant Coursework: {', '.join(edu.coursework)}")
