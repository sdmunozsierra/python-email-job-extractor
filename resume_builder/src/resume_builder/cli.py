"""CLI entry point for the resume-builder console script.

Usage::

    resume-builder --json resume.json --output tailored.docx
    resume-builder --output legacy.docx   # uses built-in sample data
"""

from __future__ import annotations

import argparse
from pathlib import Path
from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from docx import Document as DocxDocument

from .person_builder import Person
from .schema_adapter import ResumeSchemaAdapter


def build_resume(person: Person) -> "DocxDocument":
    """Build a ``.docx`` document from a :class:`Person` object.

    Returns:
        A ``python-docx`` Document ready to be saved.
    """
    from docx import Document

    from .format_education import add_education_section
    from .format_experience import add_experience_section

    doc = Document()

    # ---- Personal / Header ----
    doc.add_heading(person.name, level=0)

    contact_parts = []
    if person.email:
        contact_parts.append(person.email)
    if person.phone:
        contact_parts.append(person.phone)
    if person.location:
        contact_parts.append(person.location)
    if contact_parts:
        doc.add_paragraph(" | ".join(contact_parts))

    link_parts = []
    if person.linkedin:
        link_parts.append(person.linkedin)
    if person.github:
        link_parts.append(person.github)
    if person.portfolio:
        link_parts.append(person.portfolio)
    if link_parts:
        doc.add_paragraph(" | ".join(link_parts))

    # ---- Summary ----
    if person.summary:
        doc.add_heading("Summary", level=1)
        doc.add_paragraph(person.summary)

    # ---- Skills ----
    if person.skills:
        doc.add_heading("Technical Skills", level=1)
        # Group by category
        categories: dict[str | None, list] = {}
        for skill in person.skills:
            cat = skill.category or "General"
            categories.setdefault(cat, []).append(skill)

        for category, skills in categories.items():
            skill_strs = [str(s) for s in skills]
            doc.add_paragraph(f"{category}: {', '.join(skill_strs)}")

    if person.soft_skills:
        doc.add_heading("Soft Skills", level=1)
        doc.add_paragraph(", ".join(person.soft_skills))

    # ---- Experience ----
    add_experience_section(doc, person.experience)

    # ---- Education ----
    add_education_section(doc, person.education)

    # ---- Certifications ----
    if person.certifications:
        doc.add_heading("Certifications", level=1)
        for cert in person.certifications:
            doc.add_paragraph(str(cert), style="List Bullet")

    # ---- Projects ----
    if person.projects:
        doc.add_heading("Projects", level=1)
        for proj in person.projects:
            doc.add_heading(proj.name, level=2)
            if proj.description:
                doc.add_paragraph(proj.description)
            if proj.technologies:
                doc.add_paragraph(f"Technologies: {', '.join(proj.technologies)}")
            for hl in proj.highlights:
                doc.add_paragraph(hl, style="List Bullet")

    # ---- Languages ----
    if person.languages:
        doc.add_heading("Languages", level=1)
        doc.add_paragraph(", ".join(person.languages))

    return doc


def main() -> None:
    """CLI entry point."""
    parser = argparse.ArgumentParser(
        description="Generate a .docx resume from JSON or legacy data."
    )
    parser.add_argument(
        "--json",
        dest="json_path",
        help="Path to a JSON resume file following the schema.",
    )
    parser.add_argument(
        "--output",
        "-o",
        default="resume.docx",
        help="Output .docx file path (default: resume.docx).",
    )
    args = parser.parse_args()

    if args.json_path:
        person = ResumeSchemaAdapter.from_json_file(args.json_path)
    else:
        # Fall back to built-in sample data
        try:
            from .personal_info.sergio_david_munoz_sierra import get_person
            person = get_person()
        except ImportError:
            print("No JSON file provided and no legacy data available.")
            print("Usage: resume-builder --json resume.json --output out.docx")
            return

    doc = build_resume(person)
    doc.save(args.output)
    print(f"Resume saved to {args.output}")


if __name__ == "__main__":
    main()
