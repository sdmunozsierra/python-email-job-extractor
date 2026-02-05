"""Word-doc experience formatting helpers."""

from __future__ import annotations

from typing import TYPE_CHECKING

if TYPE_CHECKING:
    from docx import Document

from .experiece_builder import Experience


def add_experience_section(doc: "Document", experiences: list) -> None:
    """Add an Experience section to a Word document.

    Args:
        doc: python-docx Document instance.
        experiences: List of :class:`Experience` objects.
    """
    if not experiences:
        return

    doc.add_heading("Experience", level=1)

    for exp in experiences:
        # Title line
        title_text = exp.role
        if exp.company_name:
            title_text += f" -- {exp.company_name}"
        doc.add_heading(title_text, level=2)

        # Metadata line
        meta_parts = []
        if exp.location:
            meta_parts.append(exp.location)
        if exp.dates:
            meta_parts.append(exp.dates)
        if meta_parts:
            doc.add_paragraph(" | ".join(meta_parts)).italic = True

        # Description
        if exp.description:
            doc.add_paragraph(exp.description)

        # Achievements
        for achievement in exp.achievements:
            doc.add_paragraph(achievement, style="List Bullet")

        # Technologies
        if exp.technologies:
            doc.add_paragraph(f"Technologies: {', '.join(exp.technologies)}")
