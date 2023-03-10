import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def format_experience(doc, experience):

    for exp in experience:
        #p = doc.add_paragraph()
        #p.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
        text = f"{exp.role} at {exp.company_name} in {exp.location} - {exp.dates}"
        h = doc.add_heading(text, level=3)
        h.paragraph_format.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

        #p = doc.add_paragraph()
        #custom_style = doc.styles.add_style('ExperienceStyle', 1)
        #custom_style.font.name = 'Arial'
        #custom_style.font.size = Pt(12)
        #custom_style.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for proj in exp.projects:
            for j in proj:
                p.add_run(f"\n{j.name} for {j.duration}:\n").bold = True
                p.add_run(f"{j.description}\n").italic = True
                for a in j.actions:
                    p.add_run(f"- {a}\n")

def format_experience_skills(doc, experience):
    p = doc.add_paragraph()
    p.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY
    for exp in experience:
        for proj in exp.projects:
            for j in proj:
                p.add_run(f"{j.name}: ").bold = True
                p.add_run(", ".join([str(x) for x in j.skills])+"\n")